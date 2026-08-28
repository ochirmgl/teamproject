"""Document-only retrieval and Gemini answer generation for the DMS chat.

Added by: Ochir

Why this file exists:
    The teammate's original app.py already contained the DMS interface and a
    placeholder chat page, but it did not read documents or call an AI model.
    This separate module keeps the new AI/RAG logic away from the teammate's
    original interface code and makes ownership and maintenance clear.

Main responsibilities:
    - Read text from selected PDF, DOCX and TXT documents.
    - Split and locally search document text (no web search).
    - Understand Mongolian, English and common Latin-written Mongolian queries.
    - Ask Gemini to answer only from the retrieved document passages.
    - Return source labels/excerpts for the Streamlit interface.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
import math
import os
from pathlib import Path
import re
from typing import Iterable, Sequence


NO_ANSWER_MESSAGE = "Оруулсан баримт бичгүүдээс энэ талаар мэдээлэл олдсонгүй."
NO_ANSWER_MESSAGE_EN = "I could not find information about this in the selected documents."
WORD_PATTERN = re.compile(r"[0-9A-Za-zА-Яа-яЁёӨөҮү]+", re.UNICODE)
CYRILLIC_PATTERN = re.compile(r"[А-Яа-яЁёӨөҮү]")
LATIN_PATTERN = re.compile(r"[A-Za-z]")
ROMANISED_MONGOLIAN_WORDS = {
    "baigaa", "baina", "barimt", "bichig", "end", "geree", "heden",
    "medeelel", "mongol", "tuhai", "uu", "ve", "yu", "yuu",
}
ENGLISH_FOLLOW_UP_STOPWORDS = {
    "a", "about", "again", "an", "and", "any", "are", "based", "can",
    "could", "did", "do", "does", "explain", "for", "how", "i", "info",
    "information", "is", "it", "me", "more", "of", "okay", "on", "so",
    "tell", "that", "the", "then", "there", "these", "this", "those", "to",
    "us", "was", "were", "what", "why", "would", "you",
}
ENGLISH_FOLLOW_UP_REFERENCE = re.compile(
    r"\b(that|this|it|its|they|their|those|them|the above|previous answer|earlier answer)\b",
    re.IGNORECASE,
)
FOLLOW_UP_PHRASES = (
    "based on that", "according to that", "what about", "how about",
    "okay then", "so then", "and then", "tell me more", "explain more",
    "тэр талаар", "тэр мэдээлэл", "тэрний", "түүнийг", "түүний", "тэдгээр",
    "энэ талаар", "энэ нь", "энэ гэрээ", "энэ журам", "үүн дээр", "үүнээс",
    "дээрх", "өмнөх хариулт", "тэгвэл", "тэгээд", "цааш нь", "дэлгэрүүл",
)
QUERY_STOPWORDS = {
    "авч", "асуулт", "байгаа", "байна", "баримт", "баримтад", "баримтын",
    "бичсэн", "бол", "болон", "гэсэн", "гэж", "дээр", "дотор", "зүйл",
    "ийн", "ийг", "нь", "талаар", "тухай", "хэд", "хэрхэн", "энэ", "юу",
    "ямар", "яагаад", "вэ", "бэ", "уу", "үү", "юм",
}


def _question_language(question: str) -> str:
    """Choose the answer language while recognising common Mongolian transliteration."""
    words = {word.lower() for word in WORD_PATTERN.findall(question)}
    if len(words.intersection(ROMANISED_MONGOLIAN_WORDS)) >= 2:
        return "Mongolian"

    latin_count = len(LATIN_PATTERN.findall(question))
    cyrillic_count = len(CYRILLIC_PATTERN.findall(question))
    return "English" if latin_count > cyrillic_count and latin_count >= 4 else "Mongolian"


def _needs_mongolian_search_rewrite(question: str) -> bool:
    """Return True when a Latin-script query needs cross-language retrieval."""
    latin_count = len(LATIN_PATTERN.findall(question))
    cyrillic_count = len(CYRILLIC_PATTERN.findall(question))
    return latin_count > cyrillic_count and latin_count >= 4


def _needs_conversation_context(question: str) -> bool:
    """Added by Ochir: detect follow-ups that depend on an earlier turn."""
    normalised = " ".join(question.lower().split())
    if ENGLISH_FOLLOW_UP_REFERENCE.search(normalised):
        return True
    if any(phrase in normalised for phrase in FOLLOW_UP_PHRASES):
        return True

    meaningful_words = [
        word.lower()
        for word in WORD_PATTERN.findall(question)
        if len(word) >= 3
        and word.lower() not in QUERY_STOPWORDS
        and word.lower() not in ENGLISH_FOLLOW_UP_STOPWORDS
    ]
    return len(meaningful_words) < 3


class RAGError(RuntimeError):
    """Raised when the document chat cannot complete safely."""


@dataclass(frozen=True)
class DocumentChunk:
    document_id: int
    title: str
    file_name: str
    page_label: str
    text: str
    tokens: tuple[str, ...]


@dataclass(frozen=True)
class RetrievedSource:
    number: int
    document_id: int
    title: str
    file_name: str
    page_label: str
    score: float
    context_text: str
    excerpt: str


def resolve_document_path(base_dir: str | Path, stored_path: str) -> Path:
    """Resolve both Windows and POSIX relative paths from the SQLite database."""
    base = Path(base_dir).resolve()
    normalized = stored_path.replace("\\", os.sep).replace("/", os.sep)
    candidate = Path(normalized)

    if not candidate.is_absolute():
        candidate = base / candidate
    if candidate.exists():
        return candidate.resolve()

    fallback = base / "uploaded_files" / Path(normalized).name
    return fallback.resolve()


def _normalise_text(text: str) -> str:
    text = text.replace("\x00", " ")
    return re.sub(r"\s+", " ", text).strip()


def _split_text(text: str, max_chars: int = 1500, overlap: int = 220) -> list[str]:
    """Create small overlapping chunks without requiring a tokenizer package."""
    text = _normalise_text(text)
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        if end < len(text):
            boundary = max(text.rfind("\n", start, end), text.rfind(". ", start, end))
            if boundary > start + max_chars // 2:
                end = boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)

    return chunks


def _tokenise(text: str) -> tuple[str, ...]:
    """Word and character tokens work reasonably well with Mongolian suffixes."""
    words = [word.lower() for word in WORD_PATTERN.findall(text)]
    tokens = list(words)
    for word in words:
        if len(word) >= 4:
            tokens.extend(f"#{word[index:index + 3]}" for index in range(len(word) - 2))
    return tuple(tokens)


def _extract_pdf(path: Path) -> list[tuple[str, str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RAGError("pypdf суулгагдаагүй байна. requirements.txt-ийг суулгана уу.") from exc

    try:
        reader = PdfReader(str(path))
        pages: list[tuple[str, str]] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = _normalise_text(page.extract_text() or "")
            if text:
                pages.append((f"{page_number}-р хуудас", text))
        return pages
    except Exception as exc:
        raise RAGError(f"PDF файлыг уншиж чадсангүй: {path.name}") from exc


def _extract_docx(path: Path) -> list[tuple[str, str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RAGError("python-docx суулгагдаагүй байна. requirements.txt-ийг суулгана уу.") from exc

    try:
        document = Document(str(path))
        blocks: list[str] = []
        blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    blocks.append(row_text)
        text = _normalise_text("\n".join(blocks))
        return [("Word баримт", text)] if text else []
    except Exception as exc:
        raise RAGError(f"Word файлыг уншиж чадсангүй: {path.name}") from exc


def _extract_text_file(path: Path) -> list[tuple[str, str]]:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        text = _normalise_text(text)
        return [("Текст баримт", text)] if text else []
    except Exception as exc:
        raise RAGError(f"Текст файлыг уншиж чадсангүй: {path.name}") from exc


class DocumentRAG:
    """Added by Ochir: local document index plus grounded Gemini answering."""

    def __init__(self, records: Sequence[dict], base_dir: str | Path):
        self.base_dir = Path(base_dir).resolve()
        self.chunks: list[DocumentChunk] = []
        self.errors: list[str] = []
        self.document_ids: set[int] = set()
        self.document_sections: dict[int, list[tuple[str, str]]] = {}
        self.document_metadata: dict[int, dict[str, str]] = {}

        for record in records:
            document_id = int(record["id"])
            title = str(record["title"])
            stored_path = str(record["file_path"])
            path = resolve_document_path(self.base_dir, stored_path)

            if not path.exists():
                self.errors.append(f"Файл олдсонгүй: {path.name}")
                continue

            try:
                suffix = path.suffix.lower()
                if suffix == ".pdf":
                    sections = _extract_pdf(path)
                elif suffix == ".docx":
                    sections = _extract_docx(path)
                elif suffix == ".txt":
                    sections = _extract_text_file(path)
                else:
                    self.errors.append(f"AI чат дэмжихгүй файлын төрөл: {path.name}")
                    continue
            except RAGError as exc:
                self.errors.append(str(exc))
                continue

            if not sections:
                self.errors.append(
                    f"Уншигдах текст олдсонгүй: {path.name}. Энэ файл зураг хэлбэрийн PDF бол OCR шаардлагатай."
                )
                continue

            self.document_ids.add(document_id)
            self.document_sections[document_id] = sections
            self.document_metadata[document_id] = {
                "title": title,
                "file_name": path.name,
                "file_type": str(record.get("file_type") or path.suffix.lower()),
            }
            for page_label, section_text in sections:
                for chunk_text in _split_text(section_text):
                    searchable_text = f"{title}\n{path.name}\n{chunk_text}"
                    self.chunks.append(
                        DocumentChunk(
                            document_id=document_id,
                            title=title,
                            file_name=path.name,
                            page_label=page_label,
                            text=chunk_text,
                            tokens=_tokenise(searchable_text),
                        )
                    )

        self._document_frequency = Counter()
        for chunk in self.chunks:
            self._document_frequency.update(set(chunk.tokens))
        self._average_length = (
            sum(len(chunk.tokens) for chunk in self.chunks) / len(self.chunks)
            if self.chunks
            else 0.0
        )

    @property
    def document_count(self) -> int:
        return len(self.document_ids)

    @property
    def chunk_count(self) -> int:
        return len(self.chunks)

    def retrieve(
        self,
        question: str,
        selected_document_ids: Iterable[int] | None = None,
        limit: int = 3,
    ) -> list[RetrievedSource]:
        if not self.chunks:
            return []

        selected = set(int(value) for value in selected_document_ids or [])
        query_tokens = Counter(_tokenise(question))
        if not query_tokens:
            return []
        query_words = {
            token
            for token in query_tokens
            if not token.startswith("#") and len(token) >= 4 and token not in QUERY_STOPWORDS
        }

        chunk_count = len(self.chunks)
        average_length = max(self._average_length, 1.0)
        k1 = 1.5
        b = 0.75
        scored: list[tuple[float, DocumentChunk]] = []

        for chunk in self.chunks:
            if selected and chunk.document_id not in selected:
                continue
            frequencies = Counter(chunk.tokens)
            chunk_words = {token for token in frequencies if not token.startswith("#")}
            if query_words and not query_words.intersection(chunk_words):
                continue
            score = 0.0
            length_normaliser = k1 * (1 - b + b * len(chunk.tokens) / average_length)

            for token, query_frequency in query_tokens.items():
                frequency = frequencies.get(token, 0)
                if not frequency:
                    continue
                document_frequency = self._document_frequency.get(token, 0)
                inverse_frequency = math.log(
                    1 + (chunk_count - document_frequency + 0.5) / (document_frequency + 0.5)
                )
                score += (
                    inverse_frequency
                    * (frequency * (k1 + 1) / (frequency + length_normaliser))
                    * min(query_frequency, 2)
                )

            if score > 0:
                scored.append((score, chunk))

        scored.sort(key=lambda item: item[0], reverse=True)
        if scored:
            relative_cutoff = scored[0][0] * 0.35
            top_items = [item for item in scored if item[0] >= relative_cutoff][:limit]
        else:
            top_items = []

        # A selected single document can still be summarised with a very general question.
        if not top_items and len(selected) == 1:
            fallback_chunks = [chunk for chunk in self.chunks if chunk.document_id in selected][:limit]
            top_items = [(0.01, chunk) for chunk in fallback_chunks]

        grouped: dict[tuple[int, str], dict] = {}
        for score, chunk in top_items:
            key = (chunk.document_id, chunk.page_label)
            if key not in grouped:
                grouped[key] = {"score": score, "chunk": chunk, "texts": [chunk.text]}
            elif chunk.text not in grouped[key]["texts"]:
                grouped[key]["texts"].append(chunk.text)

        sources = []
        for index, item in enumerate(grouped.values(), start=1):
            chunk = item["chunk"]
            context_text = " ".join(item["texts"]).strip()
            sources.append(
                RetrievedSource(
                    number=index,
                    document_id=chunk.document_id,
                    title=chunk.title,
                    file_name=chunk.file_name,
                    page_label=chunk.page_label,
                    score=item["score"],
                    context_text=context_text,
                    excerpt=context_text[:600].strip(),
                )
            )
        return sources

    def _selected_ids(self, selected_document_ids: Iterable[int] | None) -> set[int]:
        selected = {int(value) for value in selected_document_ids or []}
        return selected or set(self.document_ids)

    def _is_summary_request(self, question: str) -> bool:
        normalised = " ".join(question.lower().split())
        patterns = (
            "summarize",
            "summary",
            "explain",
            "what is in",
            "what is it about",
            "what is this document",
            "юуны тухай",
            "тайлбарла",
            "тайлбар",
            "товчло",
            "товч",
            "агуулга",
        )
        return any(pattern in normalised for pattern in patterns)

    def _mentioned_document_ids(
        self,
        question: str,
        selected_document_ids: Iterable[int] | None,
    ) -> list[int]:
        selected = self._selected_ids(selected_document_ids)
        question_words = {
            word.lower() for word in WORD_PATTERN.findall(question) if len(word) >= 3
        }
        candidates: list[tuple[float, int]] = []

        for document_id in selected:
            metadata = self.document_metadata.get(document_id)
            if not metadata:
                continue
            title_words = {
                word.lower()
                for word in WORD_PATTERN.findall(metadata["title"])
                if len(word) >= 3 and word.lower() not in QUERY_STOPWORDS
            }
            if not title_words:
                continue
            overlap = len(question_words.intersection(title_words)) / len(title_words)
            if overlap >= 0.4:
                candidates.append((overlap, document_id))

        if not candidates and len(selected) == 1:
            return list(selected)
        if not candidates:
            return []

        candidates.sort(reverse=True)
        best_score = candidates[0][0]
        return [document_id for score, document_id in candidates if score == best_score]

    def _summary_sources(self, document_ids: Sequence[int]) -> list[RetrievedSource]:
        sources: list[RetrievedSource] = []
        total_context_chars = 0
        max_total_chars = 30000

        for document_id in document_ids:
            metadata = self.document_metadata.get(document_id)
            if not metadata:
                continue
            for page_label, section_text in self.document_sections.get(document_id, []):
                if total_context_chars >= max_total_chars:
                    break
                remaining = max_total_chars - total_context_chars
                context_text = section_text[:remaining].strip()
                if not context_text:
                    continue
                sources.append(
                    RetrievedSource(
                        number=len(sources) + 1,
                        document_id=document_id,
                        title=metadata["title"],
                        file_name=metadata["file_name"],
                        page_label=page_label,
                        score=1.0,
                        context_text=context_text,
                        excerpt=context_text[:600],
                    )
                )
                total_context_chars += len(context_text)
        return sources

    def _document_inventory(self, selected_document_ids: Iterable[int] | None) -> str:
        selected = self._selected_ids(selected_document_ids)
        lines = []
        for index, document_id in enumerate(sorted(selected), start=1):
            metadata = self.document_metadata.get(document_id)
            if metadata:
                lines.append(
                    f"{index}. {metadata['title']} | файл: {metadata['file_name']} | төрөл: {metadata['file_type']}"
                )
        return "\n".join(lines) if lines else "Сонгосон баримт байхгүй."

    def _conversation_retrieval_question(
        self,
        question: str,
        conversation_history: Sequence[dict] | None,
    ) -> str:
        """Added by Ochir: attach earlier user topics to ambiguous follow-up queries."""
        if not _needs_conversation_context(question):
            return question

        previous_questions: list[str] = []
        for item in reversed(conversation_history or []):
            if item.get("role") == "user" and item.get("content"):
                previous_questions.append(str(item["content"]))
                if len(previous_questions) == 2:
                    break

        if not previous_questions:
            return question

        previous_questions.reverse()
        previous_topic = "\n".join(previous_questions)
        return (
            f"Previous user questions that establish the topic:\n{previous_topic}\n"
            f"Current follow-up question:\n{question}"
        )

    def _rewrite_search_question(self, question: str, client, model: str) -> str:
        """Translate Latin-script questions for retrieval without answering them."""
        if not _needs_mongolian_search_rewrite(question):
            return question

        rewrite_prompt = f"""
The selected source documents are written in Mongolian. Convert the question below
into a faithful Mongolian search query for retrieving passages from legal or
organisational documents.

Rules:
- Do not answer the question.
- Do not introduce facts that are absent from the question.
- Preserve names, numbers, legal concepts and the user's intended meaning.
- Include a few close Mongolian search synonyms only when useful.
- Return only one concise Mongolian search-query line.

Question: {question}
""".strip()

        try:
            from google.genai import types

            response = client.models.generate_content(
                model=model,
                contents=rewrite_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0,
                    max_output_tokens=220,
                ),
            )
            rewritten = _normalise_text(response.text or "")
            if rewritten:
                return f"{question}\n{rewritten}"
        except Exception:
            # The normal answer call below will still provide the useful API error,
            # while Mongolian questions continue to work without this extra step.
            pass
        return question

    def answer(
        self,
        question: str,
        api_key: str,
        model: str,
        selected_document_ids: Iterable[int] | None = None,
        conversation_history: Sequence[dict] | None = None,
    ) -> tuple[str, list[RetrievedSource]]:
        try:
            from google import genai
            from google.genai import types
        except ImportError as exc:
            raise RAGError("google-genai суулгагдаагүй байна. requirements.txt-ийг суулгана уу.") from exc

        client = genai.Client(api_key=api_key)
        answer_language = _question_language(question)
        no_answer_message = (
            NO_ANSWER_MESSAGE_EN if answer_language == "English" else NO_ANSWER_MESSAGE
        )

        retrieval_question = self._conversation_retrieval_question(
            question,
            conversation_history,
        )

        retrieval_question = self._rewrite_search_question(
            retrieval_question,
            client=client,
            model=model,
        )

        mentioned_ids = self._mentioned_document_ids(retrieval_question, selected_document_ids)
        if self._is_summary_request(question) and mentioned_ids:
            sources = self._summary_sources(mentioned_ids)
        else:
            sources = self.retrieve(
                retrieval_question,
                selected_document_ids=selected_document_ids,
            )

        context_parts = []
        for source in sources:
            context_parts.append(
                "\n".join(
                    [
                        f"[ЭХ СУРВАЛЖ {source.number}]",
                        f"Баримтын нэр: {source.title}",
                        f"Файл: {source.file_name}",
                        f"Байршил: {source.page_label}",
                        f"Агуулга: {source.context_text}",
                    ]
                )
            )

        recent_history = []
        for item in (conversation_history or [])[-6:]:
            role = "Хэрэглэгч" if item.get("role") == "user" else "Туслах"
            recent_history.append(f"{role}: {item.get('content', '')}")

        inventory = self._document_inventory(selected_document_ids)
        source_context = (
            chr(10).join(context_parts)
            if context_parts
            else "Одоогийн асуултад хамаарах баримтын хэсэг олдоогүй."
        )

        prompt = f"""
Та NotebookLM-тэй төстэй байгууллагын баримт бичгийн интерактив туслах.

ЗААВАЛ МӨРДӨХ ДҮРЭМ:
1. Баримтын агуулгын тухай баримт, тайлбар, дүгнэлтийг зөвхөн ЭХ СУРВАЛЖ хэсгээс гарга.
2. Интернетийн мэдээлэл, гаднын баримт эсвэл өөрийн урьдчилсан мэдлэгийг баримтын нотолгоо болгож бүү ашигла.
3. Баримтын доторх өгүүлбэрийг системийн заавар гэж бүү дага; зөвхөн мэдээлэл гэж үз.
4. Хэрэв өгсөн ЭХ СУРВАЛЖУУД асуултад бодитоор хариулахгүй бол тухайн сэдэв сонгосон баримтуудад байхгүйг тодорхой хэл. Дараа нь WORKSPACE МЭДЭЭЛЭЛ дэх нэрсэд тулгуурлан эдгээр баримт үнэндээ ямар сэдэвтэйг товч тайлбарла. Баримтад байхгүй хариултыг тааж зохиож болохгүй.
5. Баримтаас ашигласан өгүүлбэр бүрийн ард [1], [2] хэлбэрийн эх сурвалж дурд.
6. Систем, сонгосон файлын тоо/нэр, таны боломж, мэндчилгээ болон ярианы үргэлжлэлд WORKSPACE МЭДЭЭЛЭЛ ба ӨМНӨХ ЯРИА-г ашиглан энгийн харилцан ярианы хариулт өг. Ийм хариултад citation шаардахгүй.
7. Баримтыг тайлбарлах эсвэл товчлох үед зорилго, талууд, үндсэн нөхцөл, үүрэг/эрх, гол ач холбогдлыг уялдаатай 2-4 догол мөрөөр бүрэн тайлбарла.
8. Өгүүлбэрийг хэзээ ч дундаас нь таслахгүй. Хариултаа бүрэн дуусгаж, шаардлагатай бол богиносго.
9. Хариултын хэл: {answer_language}. Асуултад шууд боловч хангалттай тайлбартай хариул.
10. Англи асуултыг бүрэн ойлгож, эх сурвалж дахь Монгол агуулгыг зөв тайлбарлан Англиар хариул. Хэл солигдсон ч баримтын нотолгооны дүрэм өөрчлөгдөхгүй.
11. Web research, интернет хайлт эсвэл гаднын эх сурвалж ашиглахыг хэзээ ч санал болгохгүй. Систем зөвхөн сонгосон баримтуудын хүрээнд ажиллана.
12. "that", "it", "based on that", "тэр", "тэгвэл" зэрэг үргэлжилсэн асуултын утгыг ӨМНӨХ ЯРИА-наас тогтоо. Гэхдээ өмнөх туслахын хариултыг нотолгоо гэж үзэхгүй; баримтын мэдээлэл бүр ЭХ СУРВАЛЖ хэсэгт дахин байх ёстой.
13. Өмнөх яриаг харсан ч асуулт яг юуг зааж байгаа нь тодорхойгүй хэвээр бол нэг богино тодруулах асуулт асуу. Шууд "мэдээлэл олдсонгүй" гэж битгий дүгнэ.

WORKSPACE МЭДЭЭЛЭЛ:
Сонгосон баримтын тоо: {len(self._selected_ids(selected_document_ids))}
{inventory}

ӨМНӨХ ЯРИА (зөвхөн одоогийн асуултын утгыг ойлгоход ашиглана; нотолгоо биш):
{chr(10).join(recent_history) if recent_history else "Өмнөх яриа байхгүй."}

ЭХ СУРВАЛЖУУД:
{source_context}

ХЭРЭГЛЭГЧИЙН АСУУЛТ:
{question}
""".strip()

        try:
            response = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.2,
                    max_output_tokens=2500,
                ),
            )
            answer_text = (response.text or "").strip()
        except Exception as exc:
            raise RAGError(f"Gemini API дуудлага амжилтгүй боллоо: {exc}") from exc

        if not answer_text:
            return no_answer_message, sources
        return answer_text, sources
